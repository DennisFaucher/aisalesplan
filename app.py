#!/usr/bin/env python3
"""
Flask web application for Perplexity AI Research
"""

import os
import re
import io
import requests
from flask import Flask, render_template, request, jsonify, send_file
import markdown

app = Flask(__name__)

# Get API key from environment variable
PERPLEXITY_API_KEY = os.getenv("PERPLEXITY_API_KEY", "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")


def search_perplexity(api_key: str, query: str):
    """
    Search using Perplexity API.
    
    Args:
        api_key: Perplexity API key
        query: Search query string
        
    Returns:
        API response as dictionary or None if error
    """
    url = "https://api.perplexity.ai/chat/completions"
    
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": "sonar",
        "messages": [
            {
                "role": "system",
                "content": "You are a helpful assistant that searches the web and provides structured information."
            },
            {
                "role": "user",
                "content": query
            }
        ],
        "temperature": 0.2,
        "max_tokens": 4000
    }
    
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=60)
        response.raise_for_status()
        return response.json()
    except requests.exceptions.RequestException as e:
        print(f"Error making API request: {e}")
        if hasattr(e, 'response') and e.response is not None:
            print(f"Response status: {e.response.status_code}")
            print(f"Response body: {e.response.text}")
        return None


def markdown_to_html_table(markdown_text: str) -> str:
    """
    Convert markdown table to HTML table.
    Also handles footnotes and URLs.
    """
    # Convert markdown to HTML using markdown library with table extension
    md = markdown.Markdown(extensions=['tables', 'fenced_code', 'nl2br'])
    html = md.convert(markdown_text)
    
    # Add some styling to make tables look better
    html = html.replace('<table>', '<table class="result-table">')
    
    # Post-process: Convert plain text table titles before tables into h2 headings
    # Pattern: Look for text like "WWT Capabilities", "WWT ATC Labs", "WWT Experts" immediately before tables
    # Find patterns where text (possibly with some whitespace) appears directly before a table
    # Common table titles to look for
    table_titles = ['WWT Capabilities', 'WWT ATC Labs', 'WWT Experts']
    
    for title in table_titles:
        # Pattern: text that might be the title (could be in a paragraph, or plain text) followed by a table
        # Try to catch cases where title is in <p> tags or plain text before <table>
        patterns = [
            # Title in paragraph tag before table
            (rf'<p>({re.escape(title)})</p>\s*<table', rf'<h2>\1</h2>\n<table'),
            # Title as plain text before table (with possible whitespace/newlines)
            (rf'({re.escape(title)})\s*\n\s*<table', rf'<h2>\1</h2>\n<table'),
            # Title with possible markdown formatting issues
            (rf'({re.escape(title)})\s+</p>\s*<table', rf'<h2>\1</h2>\n<table'),
        ]
        
        for pattern, replacement in patterns:
            html = re.sub(pattern, replacement, html, flags=re.IGNORECASE)
    
    # Post-process: Remove footnote notation (like [7], [8], etc.) from Experts section
    # Match content from "WWT Experts" heading until the next h2 heading or end of content
    def clean_experts_section(match):
        section_html = match.group(1)
        # Remove all footnote notation patterns like [1], [2], [7], [12], etc.
        # Pattern: [ followed by one or more digits followed by ]
        cleaned = re.sub(r'\[\d+\]', '', section_html)
        return cleaned
    
    # Match the entire Experts section (heading + all content until next heading or end)
    experts_section_pattern = r'(<h2[^>]*>.*?WWT\s+Experts.*?</h2>.*?)(?=<h2|</body>|</html>|$)'
    html = re.sub(experts_section_pattern, clean_experts_section, html, flags=re.DOTALL | re.IGNORECASE)
    
    return html


def _sanitize_filename_component(value: str) -> str:
    value = (value or "").strip()
    value = re.sub(r"\s+", " ", value)
    value = re.sub(r"[^A-Za-z0-9 _.-]+", "", value)
    value = value.strip().replace(" ", "_")
    return value or "export"


def _remove_experts_footnote_markers_from_markdown(markdown_text: str) -> str:
    """
    Remove bracketed numeric footnote markers (e.g. [7]) from the WWT Experts section only.
    """
    if not markdown_text:
        return markdown_text

    lines = markdown_text.splitlines()
    start_idx = None
    for i, line in enumerate(lines):
        if "wwt experts" in line.lower():
            start_idx = i
            break
    if start_idx is None:
        return markdown_text

    end_idx = len(lines)
    for j in range(start_idx + 1, len(lines)):
        low = lines[j].lower().strip()
        if low.startswith("## ") or low.startswith("# "):
            end_idx = j
            break
        if low in ("wwt capabilities", "wwt atc labs"):
            end_idx = j
            break

    before = lines[:start_idx]
    section = lines[start_idx:end_idx]
    after = lines[end_idx:]

    cleaned_section = [re.sub(r"\[\d+\]", "", l) for l in section]
    return "\n".join(before + cleaned_section + after)


def _parse_markdown_table(table_lines: list[str]) -> tuple[list[str], list[list[str]]]:
    """
    Parse a markdown pipe table into (headers, rows).
    Expects first line header, second separator, then N data rows.
    """
    def split_row(row: str) -> list[str]:
        row = row.strip()
        if row.startswith("|"):
            row = row[1:]
        if row.endswith("|"):
            row = row[:-1]
        return [c.strip() for c in row.split("|")]

    if len(table_lines) < 2:
        return [], []
    headers = split_row(table_lines[0])
    rows: list[list[str]] = []
    for line in table_lines[2:]:
        if not line.strip().startswith("|"):
            break
        rows.append(split_row(line))

    width = len(headers)
    norm_rows: list[list[str]] = []
    for r in rows:
        if len(r) < width:
            r = r + [""] * (width - len(r))
        elif len(r) > width:
            r = r[:width]
        norm_rows.append(r)
    return headers, norm_rows


def _markdown_blocks(markdown_text: str):
    """
    Yield blocks: ('heading', level, text), ('table', headers, rows), ('paragraph', text)
    Supports markdown headings and pipe tables.
    Also treats standalone table-title lines (e.g., "WWT Capabilities") immediately before a table as headings.
    """
    if not markdown_text:
        return

    table_titles = {"wwt capabilities", "wwt atc labs", "wwt experts"}
    lines = markdown_text.splitlines()
    i = 0
    n = len(lines)

    def is_table_start(idx: int) -> bool:
        if idx + 1 >= n:
            return False
        a = lines[idx].lstrip()
        b = lines[idx + 1].lstrip()
        if not a.startswith("|"):
            return False
        return bool(re.match(r"^\|?\s*:?-{3,}", b))

    while i < n:
        line = lines[i].rstrip("\n")
        if not line.strip():
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", line.strip())
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            yield ("heading", level, text)
            i += 1
            continue

        if (i + 1) < n and line.strip().lower() in table_titles and is_table_start(i + 1):
            yield ("heading", 2, line.strip())
            i += 1
            continue

        if is_table_start(i):
            table_lines = [lines[i].rstrip()]
            i += 1
            while i < n and lines[i].strip() and lines[i].lstrip().startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            headers, rows = _parse_markdown_table(table_lines)
            yield ("table", headers, rows)
            continue

        para = [line.strip()]
        i += 1
        while i < n:
            nxt = lines[i].rstrip("\n")
            if not nxt.strip():
                break
            if re.match(r"^(#{1,6})\s+", nxt.strip()):
                break
            if is_table_start(i):
                break
            para.append(nxt.strip())
            i += 1
        yield ("paragraph", " ".join(para).strip())


@app.route('/')
def index():
    """Render the main page."""
    return render_template('index.html')


@app.route('/search', methods=['POST'])
def search():
    """Handle the search request."""
    data = request.get_json()
    customer = data.get('customer', '').strip()
    theme = data.get('theme', 'AI').strip()
    
    if not customer:
        return jsonify({'error': 'Customer name is required'}), 400
    
    # Main query for Steps 1-3
    query = (
        f"Step 1: Search the public web for {customer}'s planned use of {theme}. "
        f"Extract and normalize the key {theme} themes from the web findings. "
        f"Write the summary to a markdown table titled '{customer} {theme} Research' "
        f"Do not include a date column. "
        f"Include the footnotes and footnote URLs following the markdown table. "

        f"Step 2: Map {customer}'s Planned Use of {theme} to WWT Capabilities. "
        f"Use the content in the '{customer} {theme} Research' table to map to WWT Capabilities. "
        f"Search wwt.com for content aligned to identified themes. "
        f"CRITICAL: Only include WWT capabilities that you can verify actually exist on wwt.com from your web search. "
        f"Do NOT create, invent, or guess capabilities. Only include capabilities that you can find and verify on wwt.com. "
        f"If you cannot verify a capability exists, do not include it in the table. "
        f"Evaluate and rank all findings by relevance to {customer}'s planned use, recency, and credibility. "
        f"Write the summary to a markdown table titled 'WWT Capabilities'. "
        f"CRITICAL: Do NOT include any footnote notation (like [1], [2], etc.) in the table cells. "
        f"Use a markdown heading (##) for the table title 'WWT Capabilities' before the table. "
        f"Do not include a date or a rank column. "
        f"Footnotes are not needed for this step. Do not include footnote notation in the table."
        #f"Include the footnotes and footnote URLs following the markdown table."

        #f"Step 3: Map {customer}'s Planned Use of {theme} to WWT ATC Labs. "
        f"Step 3: Map {theme} to WWT ATC Labs. "
        #f"Use the content in the '{customer} {theme} Research' table to map to WWT ATC Labs. "
        f"Search wwt.com/atc for labs related to {theme}. "
        f"CRITICAL: Only include WWT ATC Labs that you can verify actually exist on wwt.com/atc from your web search. "
        f"Do NOT create, invent, or guess labs. Only include labs that you can find and verify exist on wwt.com/atc. "
        f"If you cannot verify a lab exists, do not include it in the table. "
        #f"Evaluate and rank all findings by relevance to {customer}'s planned use, recency, and credibility. "
        f"Write the summary to a markdown table titled 'WWT ATC Labs'. "
        f"CRITICAL: Do NOT include any footnote notation (like [1], [2], etc.) in the table cells. "
        f"Use a markdown heading (##) for the table title 'WWT ATC Labs' before the table. "
        f"Do not include a date or a rank column. "
        f"Footnotes are not needed for this step. Do not include footnote notation in the table."
        #f"Include the footnotes and footnote URLs following the markdown table."
    )
    
    # Separate query for Step 4 (no other context)
    step4_query = (
        f"Step 4: List the names and titles of WWT {theme} Experts."
        f"Write the names and titles to a table titled WWT 'Experts'. "
        f"CRITICAL: Do NOT include any footnote notation (like [1], [2], etc.) in the table cells or anywhere in the Experts section. "
        f"Only include the expert name and title/role - no footnotes, no citation markers, no reference numbers."
        f"Limit the number of Experts to 10. If there are more than 10 Experts, only include the top 10 by relevance to {theme}  and credibility. "
        #f"Provide a footnote to each Expert with a link to their profile page on wwt.com."
        #f"All names listed must be found on the text of this page: https://www.wwt.com/category/ai-and-data/overview#ai-experts."
    )
    
    # Perform the main search (Steps 1-3)
    result = search_perplexity(PERPLEXITY_API_KEY, query)
    
    if not result:
        return jsonify({'error': 'Failed to get results from Perplexity API'}), 500
    
    # Extract the main response content
    if "choices" not in result or len(result["choices"]) == 0:
        return jsonify({'error': 'No results returned from API'}), 500
    
    main_content = result["choices"][0]["message"]["content"]
    
    # Perform separate Step 4 query
    step4_result = search_perplexity(PERPLEXITY_API_KEY, step4_query)
    if step4_result and "choices" in step4_result and len(step4_result["choices"]) > 0:
        step4_content = step4_result["choices"][0]["message"]["content"]
        # Combine the results
        combined_content = main_content + "\n\n" + step4_content
    else:
        # If Step 4 fails, just use main content
        combined_content = main_content
    
    # Convert markdown table to HTML
    html_content = markdown_to_html_table(combined_content)
    
    return jsonify({
        'success': True,
        'customer': customer,
        'theme': theme,
        'content': html_content,
        'markdown': combined_content
    })


@app.route('/export', methods=['POST'])
def export_docx():
    """
    Export the latest result to a Microsoft Word (.docx) document.
    Expects JSON: { customer: str, theme: str, markdown: str }
    """
    data = request.get_json() or {}
    customer = (data.get("customer") or "").strip()
    theme = (data.get("theme") or "").strip()
    markdown_text = data.get("markdown") or ""

    if not customer or not theme or not markdown_text:
        return jsonify({"error": "customer, theme, and markdown are required"}), 400

    try:
        from docx import Document
        from docx.shared import Pt
    except Exception:
        return jsonify({"error": "python-docx is not installed in the server environment"}), 500

    markdown_text = _remove_experts_footnote_markers_from_markdown(markdown_text)

    doc = Document()
    doc.add_heading(f"{customer} {theme} Research", 0)

    style = doc.styles["Normal"]
    if style and style.font and not style.font.size:
        style.font.size = Pt(11)

    for block in _markdown_blocks(markdown_text):
        kind = block[0]
        if kind == "heading":
            _, level, text = block
            level = max(1, min(int(level), 4))
            doc.add_heading(text, level=level)
        elif kind == "paragraph":
            _, text = block
            doc.add_paragraph(text)
        elif kind == "table":
            _, headers, rows = block
            if not headers:
                continue
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            for c, h in enumerate(headers):
                hdr_cells[c].text = h
                for run in hdr_cells[c].paragraphs[0].runs:
                    run.bold = True
            for r_idx, row in enumerate(rows, start=1):
                for c_idx, cell_text in enumerate(row):
                    table.rows[r_idx].cells[c_idx].text = cell_text

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"{_sanitize_filename_component(customer)}_{_sanitize_filename_component(theme)}_Research.docx"
    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)

